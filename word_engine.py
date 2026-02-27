import os
import re
import docx
from docxtpl import DocxTemplate

class GeradorWord:
    def mapear_variaveis(self, caminho_arquivo: str) -> list:
        if not os.path.exists(caminho_arquivo):
            print(f"❌ Erro: Arquivo não encontrado em {caminho_arquivo}")
            return []

        try:
            print(f"📂 Abrindo template para varredura manual: {caminho_arquivo}")
            doc = docx.Document(caminho_arquivo)
            texto_completo = ""

            # Extrai o texto mantendo a ordem dos parágrafos
            for p in doc.paragraphs:
                texto_completo += p.text + " "
            
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        texto_completo += celula.text + " "

            # Encontra as tags na ordem em que aparecem
            tags_encontradas = re.findall(r"\{\{\s*(\w+)\s*\}\}", texto_completo)
            
            # TRUQUE MÁGICO: Remove duplicatas mantendo a ordem original
            # O set() bagunça a ordem, o dict.fromkeys() mantém!
            lista_ordenada = list(dict.fromkeys(tags_encontradas))
            
            print(f"✅ Sucesso! Tags na ordem do Word: {lista_ordenada}")
            return lista_ordenada

        except Exception as e:
            print(f"💥 Falha na varredura: {e}")
            return []

    @staticmethod
    def criar_termo(caminho_modelo, caminho_salvar, dados):
        # Na hora de gerar, o docxtpl costuma funcionar bem se o arquivo existir
        from docxtpl import DocxTemplate
        doc = DocxTemplate(caminho_modelo)
        doc.render(dados)
        doc.save(caminho_salvar)