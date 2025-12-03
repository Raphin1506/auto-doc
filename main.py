import customtkinter as ctk
from tkinter import messagebox, filedialog
from tkcalendar import Calendar
from docx import Document
import re
import os
import sys
import requests
import threading
import unicodedata
import json

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue")

SETTINGS_FILE = "config.json"

# ------------------ CONFIGURAÇÕES ------------------ #

def salvar_config(caminho):
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump({"pasta_modelos": caminho}, f)

def carregar_config():
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                return json.load(f).get("pasta_modelos", "")
        except:
            return ""
    return ""

# ------------------ AMBIENTE EXECUTÁVEL ------------------ #

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# ------------------ UTILITÁRIOS ------------------ #

def remover_acentos(texto: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFD", texto) if unicodedata.category(c) != "Mn")

def formatar_cpf(cpf):
    cpf = re.sub(r'\D', '', cpf)[:11]
    if len(cpf) != 11:
        return cpf  # retorna sem formatação se não tiver 11 dígitos
    return re.sub(r'(\d{3})(\d{3})(\d{3})(\d{2})', r'\1.\2.\3-\4', cpf)

def formatar_rg(rg):
    rg = re.sub(r'\D', '', rg)[:9]
    return re.sub(r'(\d{2})(\d{3})(\d{3})(\d?)', r'\1.\2.\3-\4', rg)

# ------------------ FUNÇÃO DE UI ------------------ #

def run_on_ui(func):
    app.after(0, func)

# ------------------ BUSCA CEP ------------------ #

def buscar_cep_thread(cep_raw):
    cep = re.sub(r'\D', '', cep_raw)

    if len(cep) != 8:
        run_on_ui(lambda: messagebox.showwarning("CEP inválido", "CEP deve ter 8 dígitos."))
        run_on_ui(lambda: set_cep_loading(False))
        return

    run_on_ui(lambda: set_cep_loading(True))

    try:
        resp = requests.get(f"https://viacep.com.br/ws/{cep}/json/", timeout=6)
        resp.raise_for_status()
        data = resp.json()

        if data.get("erro"):
            run_on_ui(lambda: messagebox.showerror("Erro", "CEP não encontrado."))
            run_on_ui(lambda: set_cep_loading(False))
            return

        def preencher():
            endereco_entry.delete(0, "end")
            endereco_entry.insert(0, data.get("logradouro", ""))
            complemento_entry.delete(0, "end")
            complemento_entry.insert(0, data.get("complemento", ""))
            bairro_entry.delete(0, "end")
            bairro_entry.insert(0, data.get("bairro", ""))
            cidade_entry.delete(0, "end")
            cidade_entry.insert(0, data.get("localidade", ""))
            uf_entry.delete(0, "end")
            uf_entry.insert(0, data.get("uf", ""))
            set_cep_loading(False)

        run_on_ui(preencher)

    except Exception as e:
        run_on_ui(lambda: messagebox.showerror("Erro", str(e)))
        run_on_ui(lambda: set_cep_loading(False))

def buscar_cep_event(event=None):
    threading.Thread(target=buscar_cep_thread, args=(cep_entry.get().strip(),), daemon=True).start()

def set_cep_loading(loading: bool):
    cep_btn.configure(
        text="⏳ Buscando..." if loading else "🔎 Buscar CEP",
        state="disabled" if loading else "normal"
    )
    cep_entry.configure(state="disabled" if loading else "normal")

# ------------------ SELEÇÃO DE DATA ------------------ #

def selecionar_data():
    top = ctk.CTkToplevel(app)
    top.title("Selecionar Data")
    top.geometry("320x320")

    cal = Calendar(top, selectmode="day", date_pattern="dd/mm/yyyy")
    cal.pack(pady=20, expand=True, fill="both")

    ctk.CTkButton(
        top,
        text="Confirmar",
        command=lambda: (
            data_entry.delete(0, "end"),
            data_entry.insert(0, cal.get_date()),
            top.destroy()
        )
    ).pack(pady=10)

# ------------------ MODELOS ------------------ #

def atualizar_modelos(pasta):
    if not os.path.exists(pasta):
        modelo_combo.configure(values=["Nenhum modelo encontrado"])
        modelo_combo.set("Nenhum modelo encontrado")
        return

    modelos = [f for f in os.listdir(pasta) if f.endswith(".docx")]
    if not modelos:
        modelos = ["Nenhum modelo encontrado"]

    modelo_combo.configure(values=modelos)
    modelo_combo.set(modelos[0])  # seleciona o primeiro automaticamente

def selecionar_pasta_modelos():
    pasta = filedialog.askdirectory(title="Selecione a pasta dos modelos de termo")
    if pasta:
        salvar_config(pasta)
        atualizar_modelos(pasta)
        messagebox.showinfo("Sucesso", "Pasta definida com sucesso!")

# ------------------ GERAR DOCUMENTO ------------------ #

def gerar_documento():
    pasta_modelos = carregar_config()

    if not pasta_modelos or not os.path.exists(pasta_modelos):
        messagebox.showerror("Erro", "Nenhuma pasta de modelos configurada!")
        return

    modelo_nome = modelo_combo.get()
    caminho_modelo = os.path.join(pasta_modelos, modelo_nome)

    if not os.path.exists(caminho_modelo):
        messagebox.showerror("Erro", "Modelo não encontrado!")
        return

    nome_pessoa = nome_entry.get().strip()

    if not nome_pessoa:
        messagebox.showwarning("Erro", "Digite o nome completo.")
        return

    try:
        doc = Document(caminho_modelo)

        valores = {
            "{nome}": nome_pessoa,
            "{cpf}": formatar_cpf(cpf_entry.get()),
            "{rg}": formatar_rg(rg_entry.get()),
            "{endereco}": endereco_entry.get(),
            "{bairro}": bairro_entry.get(),
            "{cidade}": cidade_entry.get(),
            "{uf}": uf_entry.get(),
            "{data}": data_entry.get(),
            "{serie}": serie_entry.get(),
        }

        # Substitui placeholders em parágrafos e tabelas
        for p in doc.paragraphs:
            p.text = replace_placeholders(p.text, valores)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = replace_placeholders(cell.text, valores)

        nome_formatado = remover_acentos(nome_pessoa).replace(" ", "_")
        nome_arquivo = f"Termo_{nome_formatado}.docx"

        caminho_salvar = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=nome_arquivo,
            filetypes=[("Documentos Word", "*.docx")]
        )

        if caminho_salvar:
            doc.save(caminho_salvar)
            messagebox.showinfo("Sucesso", "Documento gerado!")

    except Exception as e:
        messagebox.showerror("Erro", f"{e}")

def replace_placeholders(text, valores):
    for ph, val in valores.items():
        text = text.replace(ph, val)
    return text

# ------------------ INTERFACE ------------------ #

app = ctk.CTk()
app.title("Gerador de Termo Automático - Raphael Vinicius")
app.geometry("920x980")

try:
    app.iconbitmap(resource_path("icon.ico"))
except:
    pass

frame = ctk.CTkScrollableFrame(app)
frame.pack(padx=20, pady=10, fill="both", expand=True)

def criar_campo(nome, largura=600):
    ctk.CTkLabel(frame, text=nome, font=("Arial", 15)).pack(pady=5)
    entry = ctk.CTkEntry(frame, width=largura, height=40)
    entry.pack()
    return entry

ctk.CTkButton(frame, text="📂 Selecionar Pasta de Modelos", command=selecionar_pasta_modelos).pack(pady=10)

pasta_salva = carregar_config()
if pasta_salva and os.path.exists(pasta_salva):
    modelos_iniciais = [f for f in os.listdir(pasta_salva) if f.endswith(".docx")]
    if not modelos_iniciais:
        modelos_iniciais = ["Nenhum modelo encontrado"]
else:
    modelos_iniciais = ["Nenhum modelo encontrado"]

ctk.CTkLabel(frame, text="Modelo de Termo", font=("Arial", 15)).pack()
modelo_combo = ctk.CTkOptionMenu(frame, values=modelos_iniciais)
modelo_combo.pack(pady=10)

if pasta_salva:
    atualizar_modelos(pasta_salva)

nome_entry = criar_campo("Nome completo")
cpf_entry = criar_campo("CPF")
rg_entry = criar_campo("RG")

ctk.CTkLabel(frame, text="CEP").pack()
cep_entry = ctk.CTkEntry(frame, width=200)
cep_entry.pack()
cep_entry.bind("<Return>", buscar_cep_event)

cep_btn = ctk.CTkButton(frame, text="🔎 Buscar CEP", command=buscar_cep_event)
cep_btn.pack(pady=5)

endereco_entry = criar_campo("Logradouro")
complemento_entry = criar_campo("Complemento")
bairro_entry = criar_campo("Bairro")
cidade_entry = criar_campo("Cidade")
uf_entry = criar_campo("UF")

data_entry = criar_campo("Data")
ctk.CTkButton(frame, text="📅 Escolher Data", command=selecionar_data).pack(pady=5)

serie_entry = criar_campo("Série")

ctk.CTkButton(app, text="📃 Gerar Documento", command=gerar_documento, width=260, height=50).pack(pady=15)

app.mainloop()
